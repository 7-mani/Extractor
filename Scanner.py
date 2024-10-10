import streamlit as st
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
import pdfplumber
import pytesseract
from PIL import Image
import re

def generate_pdf(invoice_data, items):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    text_y = height - 50  # Start from the top of the page

    # Write invoice data
    c.setFont("Helvetica-Bold", 14)
    for key, value in invoice_data.items():
        c.drawString(50, text_y, f"{key}: {value}")
        text_y -= 20  # Move down for the next line

    # Write invoice items
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, text_y, "Invoice Items:")
    text_y -= 20

    c.setFont("Helvetica", 12)
    for item in items:
        c.drawString(50, text_y, item)
        text_y -= 20  # Move down for the next line

    c.save()
    buffer.seek(0)
    return buffer

def extract_invoice_data_from_text(text):
    invoice_data = {}
    
    # Use regex to find all lines with key-value pairs
    lines = text.splitlines()
    for line in lines:
        match = re.match(r'([^\:]+):\s*(.*)', line.strip(), re.IGNORECASE)
        if match:
            key = match.group(1).strip()
            value = match.group(2).strip()
            invoice_data[key] = value

    # Set default values for specific fields if they are missing
    for key in ["Invoice Number", "Sold By", "PAN No", "GST Registration No", "CIN No", 
                "Order Number", "Order Date", "Billing Address", "Shipping Address", "Invoice Date"]:
        if key not in invoice_data:
            invoice_data[key] = "N/A"
    
    return invoice_data

def extract_invoice_data_from_pdf(pdf_text):
    return extract_invoice_data_from_text(pdf_text)

def extract_invoice_data_from_word(doc):
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    return extract_invoice_data_from_text(full_text)

def extract_text_from_image(image):
    text = pytesseract.image_to_string(image)
    return extract_invoice_data_from_text(text)

st.title("Invoice Extractor")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "png", "jpg", "jpeg"])

if uploaded_file is not None:
    invoice_data = {}
    items = []

    try:
        if uploaded_file.type == "application/pdf":
            with pdfplumber.open(uploaded_file) as pdf:
                full_text = ""
                for page in pdf.pages:
                    full_text += page.extract_text() + "\n"
                
                invoice_data = extract_invoice_data_from_pdf(full_text)

                # Extract tables if they exist
                tables = []
                for page in pdf.pages:
                    tables.extend(page.extract_tables())
                if tables:
                    for table in tables:
                        for row in table[1:]:  # Skip header row
                            if row and len(row) > 0:  # Check if the row is valid
                                items.append(", ".join(filter(None, row)))  # Join non-empty cells

        elif uploaded_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document","application/msword"]:
            doc = Document(uploaded_file)
            invoice_data = extract_invoice_data_from_word(doc)

            # Extract tables from Word document
            for table in doc.tables:
                for row in table.rows:
                    if row.cells:  # Check if the row has cells
                        items.append(", ".join(filter(None, [cell.text for cell in row.cells])))  # Join non-empty cells

        elif uploaded_file.type in ["image/png", "image/jpeg", "image/jpg"]:
            image = Image.open(uploaded_file)
            invoice_data = extract_text_from_image(image)

    except Exception as e:
        st.error(f"An error occurred: {str(e)}")

    # Display extracted data
    st.subheader("Extracted Invoice Details")
    for key, value in invoice_data.items():
        st.write(f"**{key}:** {value}")

    st.subheader("Invoice Items")
    for index, item in enumerate(items, start=1):
        st.write(f"**Item {index}:** {item}")

    # Generate PDF for download
    pdf_output = generate_pdf(invoice_data, items)
    st.download_button("Download PDF", pdf_output, "invoice_output.pdf")

else:
    st.warning("Please upload a file.")
